"""
Professional ATS-Optimized CV Generator
=======================================
Creates clean, single-column, recruiter-friendly CVs.

Design principles:
- Single column layout (ATS-friendly)
- No icons, tables, graphics, or colors
- Clean section headers with subtle separators
- Consistent typography hierarchy
- Optimal whitespace for readability
- Professional, not AI-generated feel
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether, PageBreak
from reportlab.lib.colors import HexColor, black
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from app.models import ParsedCV
import re


def format_skills_for_display(skills: list[str]) -> str:
    """
    Format skills into categorized lines for better readability.
    Attempts to group by common categories.
    """
    if not skills:
        return ""
    
    # Define category keywords
    categories = {
        "Languages": ["python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php", "swift", "kotlin", "scala", "r", "c/c++", "c", "perl", "bash", "shell"],
        "Frontend": ["react", "vue", "angular", "next.js", "nuxt", "svelte", "html", "css", "tailwind", "bootstrap", "sass", "redux", "react native", "flutter", "material ui", "chakra"],
        "Backend": ["node.js", "express", "django", "flask", "fastapi", "spring", "spring boot", ".net", "rails", "laravel", "nest.js", "express.js", "asp.net", "gin", "fiber"],
        "Databases": ["postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sqlite", "oracle", "sql server", "dynamodb", "cassandra", "firebase", "supabase", "neo4j", "mariadb"],
        "DevOps & Cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "github actions", "gitlab", "ci/cd", "terraform", "ansible", "linux", "nginx", "apache", "cloudflare", "vercel", "heroku", "digitalocean"],
        "Testing": ["jest", "pytest", "selenium", "cypress", "junit", "mocha", "testing", "unit testing", "integration testing", "e2e", "playwright", "vitest", "rspec"],
        "AI/ML": ["tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "machine learning", "deep learning", "nlp", "computer vision", "keras", "opencv", "huggingface", "langchain", "openai"],
        "Tools": ["git", "jira", "postman", "figma", "webpack", "vite", "npm", "yarn", "agile", "scrum", "rest", "graphql", "websockets", "swagger", "confluence", "slack", "notion"]
    }
    
    categorized = {cat: [] for cat in categories}
    uncategorized = []
    
    def matches(skill_lower: str, kw: str) -> bool:
        """Robust matching: avoid substring matches for very short keywords."""
        kw_l = kw.lower().strip()
        if not kw_l:
            return False
        # Exact match
        if skill_lower == kw_l:
            return True
        # Multi-word keyword or contains punctuation -> allow substring match
        if len(kw_l) >= 4 and (" " in kw_l or "." in kw_l or "/" in kw_l or "-" in kw_l):
            return kw_l in skill_lower
        # For short keywords (<=2), only match as a whole token
        if len(kw_l) <= 2:
            return any(tok == kw_l for tok in re.findall(r"[a-z0-9\+/#\.]+", skill_lower))
        # Default: word-boundary match
        return re.search(rf"\b{re.escape(kw_l)}\b", skill_lower) is not None

    for skill in skills:
        if not skill:
            continue
        skill_lower = skill.lower().strip()
        found = False
        for cat, keywords in categories.items():
            if any(matches(skill_lower, kw) for kw in keywords):
                categorized[cat].append(skill)
                found = True
                break
        if not found:
            uncategorized.append(skill)
    
    # Build output
    lines = []
    for cat, cat_skills in categorized.items():
        if cat_skills:
            lines.append(f"{cat}: {', '.join(cat_skills)}")
    
    if uncategorized:
        lines.append(f"Other: {', '.join(uncategorized)}")
    
    # If categorization didn't work well, just return comma-separated
    if len(lines) <= 1:
        return ", ".join([s for s in skills if s])
    
    return "\n".join(lines)


def create_cv_docx(cv: ParsedCV) -> bytes:
    """
    Create a professional ATS-optimized DOCX CV.
    
    Features:
    - Clean single-column layout
    - Professional typography
    - Consistent spacing
    - No graphics or tables
    """
    doc = Document()
    
    # Set narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # ─────────────────────────────────────────────────────────────
    # HEADER: Name & Contact
    # ─────────────────────────────────────────────────────────────
    if cv.personal_info.name:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(cv.personal_info.name.upper())
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.space_after = Pt(4)
    
    # Contact line
    contact_parts = []
    if cv.personal_info.email:
        contact_parts.append(cv.personal_info.email)
    if cv.personal_info.phone:
        contact_parts.append(cv.personal_info.phone)
    if cv.personal_info.location:
        contact_parts.append(cv.personal_info.location)
    if cv.personal_info.linkedin:
        # Clean LinkedIn URL
        linkedin = cv.personal_info.linkedin.replace("https://", "").replace("www.", "")
        contact_parts.append(linkedin)
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(" • ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(80, 80, 80)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.space_after = Pt(12)
    
    def add_section_header(title: str):
        """Add a clean section header with underline."""
        para = doc.add_paragraph()
        para.space_before = Pt(14)
        para.space_after = Pt(6)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Add bottom border
        para.paragraph_format.space_after = Pt(2)
        
        # Separator line
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(8)
        sep_run = sep.add_run("-" * 80)
        sep_run.font.size = Pt(6)
        sep_run.font.color.rgb = RGBColor(180, 180, 180)
    
    def add_bullet_point(text: str):
        """Add a clean bullet point."""
        # Use Word's native bullet list style for better compatibility than a literal "•".
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        run.font.size = Pt(10)
    
    # ─────────────────────────────────────────────────────────────
    # PROFESSIONAL SUMMARY
    # ─────────────────────────────────────────────────────────────
    if cv.summary:
        add_section_header("Professional Summary")
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(cv.summary)
        summary_run.font.size = Pt(10)
        summary_para.paragraph_format.space_after = Pt(6)
    
    # ─────────────────────────────────────────────────────────────
    # PROFESSIONAL EXPERIENCE
    # ─────────────────────────────────────────────────────────────
    if cv.experience:
        add_section_header("Professional Experience")
        
        for i, exp in enumerate(cv.experience):
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            company_run = title_para.add_run(f" | {exp.company}")
            company_run.font.size = Pt(11)
            title_para.paragraph_format.space_after = Pt(2)
            
            # Dates
            if exp.start_date or exp.end_date:
                date_para = doc.add_paragraph()
                date_text = f"{exp.start_date or ''} – {exp.end_date or 'Present'}"
                date_run = date_para.add_run(date_text)
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = RGBColor(100, 100, 100)
                date_run.italic = True
                date_para.paragraph_format.space_after = Pt(4)
            
            # Bullet points (combine responsibilities and achievements)
            bullets = []
            if exp.achievements:
                bullets.extend(exp.achievements)
            if exp.responsibilities:
                bullets.extend(exp.responsibilities)
            
            for bullet in bullets[:6]:  # Limit to 6 bullets per job
                if bullet and bullet.strip():
                    add_bullet_point(bullet)
            
            # Space between jobs
            if i < len(cv.experience) - 1:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(8)
    
    # ─────────────────────────────────────────────────────────────
    # EDUCATION
    # ─────────────────────────────────────────────────────────────
    if cv.education:
        add_section_header("Education")
        
        for edu in cv.education:
            edu_para = doc.add_paragraph()
            
            # Degree
            degree_text = edu.degree
            if edu.field:
                degree_text += f" in {edu.field}"
            degree_run = edu_para.add_run(degree_text)
            degree_run.bold = True
            degree_run.font.size = Pt(10)
            
            # Institution
            inst_run = edu_para.add_run(f" | {edu.institution}")
            inst_run.font.size = Pt(10)
            
            # Date
            if edu.graduation_date:
                date_run = edu_para.add_run(f" ({edu.graduation_date})")
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = RGBColor(100, 100, 100)
            
            edu_para.paragraph_format.space_after = Pt(4)
    
    # ─────────────────────────────────────────────────────────────
    # SKILLS
    # ─────────────────────────────────────────────────────────────
    if cv.skills:
        add_section_header("Technical Skills")
        # Format skills in a readable way with line breaks between categories
        skills_text = format_skills_for_display(cv.skills)
        for line in skills_text.split('\n'):
            if line.strip():
                skills_para = doc.add_paragraph()
                skills_run = skills_para.add_run(line)
                skills_run.font.size = Pt(10)
                skills_para.paragraph_format.space_after = Pt(2)
    
    # ─────────────────────────────────────────────────────────────
    # CERTIFICATIONS
    # ─────────────────────────────────────────────────────────────
    if cv.certifications:
        add_section_header("Certifications")
        for cert in cv.certifications:
            if cert and cert.strip():
                add_bullet_point(cert)
    
    # ─────────────────────────────────────────────────────────────
    # LANGUAGES
    # ─────────────────────────────────────────────────────────────
    if cv.languages:
        add_section_header("Languages")
        langs_para = doc.add_paragraph()
        langs_text = " • ".join([l for l in cv.languages if l])
        langs_run = langs_para.add_run(langs_text)
        langs_run.font.size = Pt(10)
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_cv_pdf(cv: ParsedCV) -> bytes:
    """
    Create a professional ATS-optimized PDF CV.
    
    Features:
    - Clean single-column layout
    - Professional typography (no fancy fonts)
    - Consistent spacing
    - No graphics, icons, or colors
    - Optimized for ATS parsing
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        topMargin=0.5*inch, 
        bottomMargin=0.5*inch, 
        leftMargin=0.6*inch, 
        rightMargin=0.6*inch
    )
    
    styles = getSampleStyleSheet()
    
    # ─────────────────────────────────────────────────────────────
    # CUSTOM STYLES
    # ─────────────────────────────────────────────────────────────
    
    # Name - centered, bold, larger
    styles.add(ParagraphStyle(
        name='CVName',
        fontName='Helvetica-Bold',
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=4,
        textColor=black
    ))
    
    # Contact info - centered, smaller, gray
    styles.add(ParagraphStyle(
        name='CVContact',
        fontName='Helvetica',
        fontSize=10,
        alignment=TA_CENTER,
        textColor=HexColor('#505050'),
        spaceAfter=16
    ))
    
    # Section header - bold, uppercase
    styles.add(ParagraphStyle(
        name='CVSection',
        fontName='Helvetica-Bold',
        fontSize=11,
        spaceBefore=14,
        spaceAfter=2,
        textColor=black
    ))
    
    # Section separator
    styles.add(ParagraphStyle(
        name='CVSeparator',
        fontName='Helvetica',
        fontSize=6,
        textColor=HexColor('#b4b4b4'),
        spaceAfter=8
    ))
    
    # Job title - bold
    styles.add(ParagraphStyle(
        name='CVJobTitle',
        fontName='Helvetica-Bold',
        fontSize=11,
        spaceAfter=2,
        textColor=black
    ))
    
    # Date - italic, gray
    styles.add(ParagraphStyle(
        name='CVDate',
        fontName='Helvetica-Oblique',
        fontSize=9,
        textColor=HexColor('#646464'),
        spaceAfter=4
    ))
    
    # Bullet point
    styles.add(ParagraphStyle(
        name='CVBullet',
        fontName='Helvetica',
        fontSize=10,
        leftIndent=12,
        spaceAfter=3,
        textColor=black
    ))
    
    # Normal text
    styles.add(ParagraphStyle(
        name='CVText',
        fontName='Helvetica',
        fontSize=10,
        spaceAfter=6,
        textColor=black,
        alignment=TA_JUSTIFY
    ))
    
    # Skills text
    styles.add(ParagraphStyle(
        name='CVSkills',
        fontName='Helvetica',
        fontSize=10,
        spaceAfter=6,
        textColor=black
    ))
    
    story = []
    
    def safe_text(text):
        """Escape XML special characters."""
        if not text:
            return ""
        return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    def add_section(title: str):
        """Add a section header with separator."""
        story.append(Paragraph(title.upper(), styles['CVSection']))
        story.append(Paragraph("-" * 95, styles['CVSeparator']))
    
    # ─────────────────────────────────────────────────────────────
    # HEADER
    # ─────────────────────────────────────────────────────────────
    if cv.personal_info.name:
        story.append(Paragraph(safe_text(cv.personal_info.name.upper()), styles['CVName']))
    
    contact_parts = []
    if cv.personal_info.email:
        contact_parts.append(cv.personal_info.email)
    if cv.personal_info.phone:
        contact_parts.append(cv.personal_info.phone)
    if cv.personal_info.location:
        contact_parts.append(cv.personal_info.location)
    if cv.personal_info.linkedin:
        linkedin = cv.personal_info.linkedin.replace("https://", "").replace("www.", "")
        contact_parts.append(linkedin)
    
    if contact_parts:
        story.append(Paragraph(safe_text(" • ".join(contact_parts)), styles['CVContact']))
    
    # ─────────────────────────────────────────────────────────────
    # PROFESSIONAL SUMMARY
    # ─────────────────────────────────────────────────────────────
    if cv.summary:
        add_section("Professional Summary")
        story.append(Paragraph(safe_text(cv.summary), styles['CVText']))
    
    # ─────────────────────────────────────────────────────────────
    # PROFESSIONAL EXPERIENCE
    # ─────────────────────────────────────────────────────────────
    if cv.experience:
        add_section("Professional Experience")
        
        for i, exp in enumerate(cv.experience):
            # Group each job entry together to prevent awkward page breaks
            job_elements = []
            
            # Title and company
            title_text = f"<b>{safe_text(exp.title)}</b> | {safe_text(exp.company)}"
            job_elements.append(Paragraph(title_text, styles['CVJobTitle']))
            
            # Dates
            if exp.start_date or exp.end_date:
                date_text = f"{exp.start_date or ''} – {exp.end_date or 'Present'}"
                job_elements.append(Paragraph(safe_text(date_text), styles['CVDate']))
            
            # Bullets (achievements first, then responsibilities)
            bullets = []
            if exp.achievements:
                bullets.extend(exp.achievements)
            if exp.responsibilities:
                bullets.extend(exp.responsibilities)
            
            for bullet in bullets[:6]:
                if bullet and bullet.strip():
                    job_elements.append(Paragraph(f"• {safe_text(bullet)}", styles['CVBullet']))
            
            # Use KeepTogether to prevent job entries from being split across pages
            # But only if the job has 4 or fewer bullets (otherwise it might be too long)
            if len(bullets) <= 4:
                story.append(KeepTogether(job_elements))
            else:
                # For longer entries, keep at least title + date + first 2 bullets together
                story.append(KeepTogether(job_elements[:4]))
                for elem in job_elements[4:]:
                    story.append(elem)
            
            # Space between jobs
            if i < len(cv.experience) - 1:
                story.append(Spacer(1, 10))
    
    # ─────────────────────────────────────────────────────────────
    # EDUCATION
    # ─────────────────────────────────────────────────────────────
    if cv.education:
        add_section("Education")
        
        edu_elements = []
        for edu in cv.education:
            degree = safe_text(edu.degree)
            if edu.field:
                degree += f" in {safe_text(edu.field)}"
            
            edu_text = f"<b>{degree}</b> | {safe_text(edu.institution)}"
            if edu.graduation_date:
                edu_text += f" ({safe_text(edu.graduation_date)})"
            
            edu_elements.append(Paragraph(edu_text, styles['CVText']))
        
        # Keep all education entries together if possible
        if len(edu_elements) <= 3:
            story.append(KeepTogether(edu_elements))
        else:
            for elem in edu_elements:
                story.append(elem)
    
    # ─────────────────────────────────────────────────────────────
    # TECHNICAL SKILLS
    # ─────────────────────────────────────────────────────────────
    if cv.skills:
        add_section("Technical Skills")
        # Format skills with categories
        skills_formatted = format_skills_for_display(cv.skills)
        for line in skills_formatted.split('\n'):
            if line.strip():
                story.append(Paragraph(safe_text(line), styles['CVSkills']))
    
    # ─────────────────────────────────────────────────────────────
    # CERTIFICATIONS
    # ─────────────────────────────────────────────────────────────
    if cv.certifications:
        add_section("Certifications")
        for cert in cv.certifications:
            if cert and cert.strip():
                story.append(Paragraph(f"• {safe_text(cert)}", styles['CVBullet']))
    
    # ─────────────────────────────────────────────────────────────
    # LANGUAGES
    # ─────────────────────────────────────────────────────────────
    if cv.languages:
        add_section("Languages")
        langs_text = " • ".join([safe_text(l) for l in cv.languages if l])
        story.append(Paragraph(langs_text, styles['CVSkills']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
